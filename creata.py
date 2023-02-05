import shutil
from docx import Document
import pandas as pd
from docx2pdf import convert
import os
import subprocess

def doc2pdf_linux(doc,dirs):
    """
    convert a doc/docx document to pdf format (linux only, requires libreoffice)
    :param doc: path to document
    """
    cmd = 'libreoffice --convert-to pdf'.split() + [doc] + '--outdir'.split() + [dirs]
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=10)
    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)

def main1(name='Name of Proff'):
    os.mkdir('output')
    template_file_path = './input.docx'
    os.mkdir(f'output/{name}')
    output_file_path = 'output/'+name + '/' + name + '.docx'

    variables = {
        "<name>": name,
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(
                            paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    data = pd.read_csv('index.csv')
    data = data.to_numpy()
    info = [i[0] for i in data]
    for i in info:
        main1(name=i)
        doc2pdf_linux(f'output/{i}/{i}.docx',f'output/{i}/')
    shutil.make_archive('output','zip','output')